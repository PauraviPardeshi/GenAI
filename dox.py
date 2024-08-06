import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Use the raw string to specify the full path to the file
file_name = r"C:\Users\Pauravi\Technology\profile.txt"

with open(file_name, 'r', encoding='utf-8') as file:
    response_profile = {"response": file.read()}

response_techprofile = response_profile['response']

# Technologiename extrahieren
technology_name_match = re.search(r'### Technologiename:\s*(.*)', response_techprofile)
if technology_name_match:
    technology_name = technology_name_match.group(1).strip()
else:
    technology_name = "Unbekannte_Technologie"

# Speichern des Inhalts in einer .docx Datei
current_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
file_path = rf"C:\Users\Pauravi\Technology\{technology_name.replace(' ', '_')}_{current_time}.docx"

# Erstellen eines neuen Dokuments
doc = Document()
doc.add_heading(technology_name, 0)

# Parsing and formatting the content
lines = response_techprofile.split('\n')
for line in lines:
    if line.startswith('###'):
        # Extract the text after ###
        section_title = line[3:].strip()
        if section_title:
            # Add section title in bold
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(section_title)
            run.bold = True
            paragraph.style = doc.styles['Normal']
    else:
        # Add the rest of the lines as normal text
        doc.add_paragraph(line)

# Speichern des Dokuments
doc.save(file_path)

print(f"Der Steckbrief '{technology_name}' wurde erfolgreich in '{file_path}' gespeichert.")
